from docx import Document
template_path = "/sample.docx"

def load_template(template_path):
    return Document(template_path)

def replace_placeholders(doc, data):
    for paragraph in doc.paragraphs:
        print(f"og text: {paragraph.text}")

        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                print(f"found placeholder: {placeholder}")
                print(f"replaving {placeholder} with {value}")
                paragraph.text = paragraph.text.replace(placeholder, str(value))
                print(f"updated text: {paragraph.text}")


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in paragraph.text:
                            print(f"replacing {placeholder} with {value}")
                            paragraph.text = paragraph.text.replace(placeholder, str(value))
    return doc

def save_document(doc, output_path):
    doc.save(output_path)

def generate_resume(template_path, output_path, data):
    doc = load_template(template_path)
    doc = replace_placeholders(doc, data)
    save_document(doc, output_path)
    print(f" saved to {output_path}")

# simple data
# data = {
#     "First_Name": "John",
#     "Last_Name": "Doe",
#     "Phone_Number": "+1234567890",
#     "Email": "john.doe@example.com",
#     "LinkedIn": "linkedin.com/in/johndoe",
#     "Skill_1": "Python",
#     "Skill_2": "Data Analysis",
#     "Job_Title": "Software Engineer",
#     "Company_Name": "Tech Corp",
#     "Start_Date": "Jan 2020",
#     "START_DATE": "Jan 2020",
#     "Year": "Jan 2020",
#     "YEAR": "Jan 2020",
#     "End_Date": "Present",
#     "END_DATE": "Present",
#     "Location": "San Francisco, CA",
#     "Description": "Developed scalable backend systems.",
#     "Degree": "Bachelor of Science in Computer Science",
#     "Institution": "University of Michigan",
#     "Graduation_Date": "May 2019"
# }


data = {
    "First_Name": "John",
    "Last_Name": "Doe",
    "Phone_Number": "+1 (234) 567-8901",
    "Email": "john.doe@example.com",
    "LinkedIn": "linkedin.com/in/johndoe",
    "GitHub": "github.com/johndoe",
    "Portfolio": "johndoe.dev",
    "Address": "1234 Elm Street, San Francisco, CA, 94107",
    "Summary": "Experienced Software Engineer with 5+ years in full-stack development, specializing in scalable web applications and cloud-based solutions.",

    "Skill_1a": "Python",
    "Skill_1b": "JavaScript",
    "Skill_1c": "Java",
    "Skill_1d": "C++",
    "Skill_2a": "Django",
    "Skill_2b": "React.js",
    "Skill_2c": "Node.js",
    "Skill_2d": "Spring Boot",
    "Skill_3a": "PostgreSQL",
    "Skill_3b": "MongoDB",
    "Skill_3c": "MySQL",
    "Skill_4a": "AWS",
    "Skill_4b": "Azure",
    "Skill_4c": "Docker",
    "Skill_4d": "Kubernetes",
    "Skill_5a": "Git",
    "Skill_5b": "CI/CD",
    "Skill_5c": "Agile Methodologies",

    "Job_Title_1": "Senior Software Engineer",
    "Company_Name_1": "Tech Corp",
    "Start_Date_1": "Jan 2021",
    "END_DATE_1": "Present",
    "Location_1": "San Francisco, CA",
    "Description_1a": "Designed and implemented microservices architecture for scalable applications.",
    "Description_1b": "Led a team of 5 engineers to develop AI-powered recommendation systems.",
    "Description_1c": "Optimized database queries, reducing query times by 40%.",

    "Job_Title_2": "Software Engineer",
    "Company_Name_2": "Innovate Inc.",
    "Start_Date_2": "Jun 2018",
    "END_DATE_2": "Dec 2020",
    "Location_2": "Seattle, WA",
    "Description_2a": "Developed RESTful APIs for enterprise applications.",
    "Description_2b": "Integrated third-party authentication services for enhanced security.",
    "Description_2c": "Collaborated with UI/UX designers to improve frontend performance.",

    "Degree_1": "Master of Science in Computer Science",
    "Institution_1": "Stanford University",
    "Graduation_Date_1": "May 2018",

    "Degree_2": "Bachelor of Science in Computer Science",
    "Institution_2": "University of Michigan",
    "Graduation_Date_2": "May 2016",

    "Certification_1": "AWS Certified Solutions Architect",
    "Issuing_Organization_1": "Amazon Web Services",
    "Issue_Date_1": "Sep 2021",
    "Expiration_Date_1": "Sep 2024",

    "Certification_2": "Google Professional Cloud Developer",
    "Issuing_Organization_2": "Google Cloud",
    "Issue_Date_2": "Mar 2020",
    "Expiration_Date_2": "Mar 2023",

    "Project_1": "SmartHome AI",
    "Project_Description_1": "Developed an AI-powered home automation system using Python and IoT technologies.",
    "Project_Technologies_1a": "Python",
    "Project_Technologies_1b": "TensorFlow",
    "Project_Technologies_1c": "Raspberry Pi",
    "Project_GitHub_1": "github.com/johndoe/smarthome-ai",

    "Project_2": "E-commerce Platform",
    "Project_Description_2": "Built a scalable e-commerce platform using Django and React.js.",
    "Project_Technologies_2a": "Django",
    "Project_Technologies_2b": "React.js",
    "Project_Technologies_2c": "PostgreSQL",
    "Project_GitHub_2": "github.com/johndoe/ecommerce-platform",

    "Publication_1": "Optimizing Machine Learning Pipelines for Scalable Deployment",
    "Publication_Journal_1": "IEEE Transactions on Software Engineering",
    "Publication_Date_1": "July 2022",
    "Publication_DOI_1": "10.1109/TSE.2022.1234567",

    "Language_1": "English",
    "Proficiency_1": "Fluent",
    "Language_2": "Spanish",
    "Proficiency_2": "Intermediate",

    "Volunteer_Role_1": "Mentor",
    "Volunteer_Organization_1": "Girls Who Code",
    "Volunteer_Start_Date_1": "Jan 2020",
    "Volunteer_END_DATE_1": "Present",
    "Volunteer_Location_1": "Remote",
    "Volunteer_Description_1": "Mentored high school students in coding and career development.",

    "Award_1": "Best AI Startup Pitch",
    "Award_Issuing_Organization_1": "Tech Innovators Summit",
    "Award_Issue_Date_1": "Oct 2023"
}




generate_resume(template_path, "generated_resume0.docx", data)




# from docx import Document
#
# # Create a new document
# doc = Document()
#
# # Add a paragraph
# doc.add_paragraph("Hello, this is a sample document!")
#
# # Save the document
# doc.save("sample.docx")